"""
tailor.py
---------
Takes high-match jobs from search_agent.py, uses Claude to tailor
Stella's resume for each one, and saves a .docx file per job.

Called by search_agent.py — import run() directly, or run standalone.

Requirements:
    pip install anthropic python-docx python-dotenv

Environment variables (set in GitHub Secrets):
    ANTHROPIC_API_KEY
"""

import os
import json
import re
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from datetime import datetime

load_dotenv()

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

MATCH_THRESHOLD = 80  # Only tailor resumes for jobs scoring above this

STELLA_RESUME = {
    "name": "STELLA FUENTES",
    "contact": "Vancouver, BC  •  778-792-0815  •  stella_fuentes@outlook.com  •  linkedin.com/in/stellafuentes  •  github.com/stellaxf8",
    "summary": "Business Analyst and AI specialist with expertise in scoping and delivering AI-powered automation solutions. Experienced in requirements gathering, user story development, and process mapping. Azure AI Engineer certified with hands-on experience across LLM APIs and Python.",
    "skills": {
        "Business Analysis": "Requirements gathering, user stories, process mapping, stakeholder coordination",
        "Programming": "Python, SQL, Node.js",
        "AI / ML": "LLM fine-tuning, prompt engineering, OpenAI API, Anthropic API, RAG",
        "Tools": "Azure AI, Power BI, Replit, Power Automate, Pandas, Plotly",
    },
    "experience": [
        {
            "title": "AI Analyst / Business Analyst",
            "company": "Red Pill Labs",
            "location": "Vancouver, BC",
            "dates": "Jun 2025 – Present",
            "bullets": [
                "Owned end-to-end delivery of RFP analysis tool (Replit/Claude), reducing vendor evaluation from 3 hours to 30 minutes; in active organizational use.",
                "Led discovery and requirements gathering for AI knowledge management tool; developed user stories and functional specifications.",
                "Built Python automation solutions (Pandas, Scikit-learn) reducing manual processing by 35%.",
                "Created Power BI dashboards analyzing 50K+ records and defined KPIs informing quarterly decisions.",
            ],
        },
        {
            "title": "Manager / Bartender",
            "company": "Earls Kitchen + Bar",
            "location": "Surrey, BC",
            "dates": "Sep 2017 – Jan 2022",
            "bullets": [
                "Managed 15-person team in high-volume environment; drove data-driven operational decision-making.",
            ],
        },
    ],
    "projects": [
        {
            "title": "RFP Analysis & Vendor Intelligence Platform",
            "date": "Jan 2026",
            "description": "Built AI-powered tool in Replit using Claude and Python; engineered prompts for executive summaries and integrated web search. Reduced turnaround by 60%; in active use.",
        },
        {
            "title": "Predictive Finance for Credit Union — AI Model",
            "date": "Apr 2025",
            "description": "Gathered requirements and built AI forecasting model to help credit unions manage daily cash positions by analyzing transaction patterns. In active use.",
        },
    ],
    "education": [
        "Diploma, Business Information Technology Management (AI Management Option) | BCIT | Sep 2023 – May 2025",
        "Microsoft Certified: Azure AI Engineer Associate — May 2026",
        "Microsoft Applied Skills: Streamline business workflows with AI chat — Jun 2026",
        "Developing LLM Applications with LangChain — DataCamp, Jun 2026",
        "Talking to AI: Prompt Engineering for Project Managers — PMI, Aug 2025",
    ],
}


# ── Claude Tailoring ───────────────────────────────────────────────────────────

def tailor_with_claude(job: dict, client: anthropic.Anthropic) -> dict:
    """
    Sends the resume + job description to Claude and gets back a tailored
    version as structured JSON.
    """
    resume_text = json.dumps(STELLA_RESUME, indent=2)

    prompt = f"""You are a professional resume writer. Tailor this resume for the job posting below.

RESUME (JSON):
{resume_text}

JOB POSTING:
Title: {job['title']}
Company: {job['company']}
Description: {job['description']}

Rules:
- Rewrite the summary to mirror the job's language and priorities (2-3 sentences max).
- Reorder the experience bullets so the most relevant ones appear first.
- You may lightly rephrase bullets to include keywords from the JD, but do not invent new experience.
- Reorder the skills dict so the most relevant categories appear first.
- Do not remove any jobs or projects.
- Keep all dates, companies, and titles exactly as-is.

Return ONLY a JSON object with this exact structure (no preamble, no markdown):
{{
  "summary": "<tailored summary>",
  "skills": {{
    "<category>": "<skills string>",
    ...
  }},
  "experience": [
    {{
      "title": "<same>",
      "company": "<same>",
      "location": "<same>",
      "dates": "<same>",
      "bullets": ["<reordered/rephrased bullets>"]
    }}
  ],
  "changes_made": "<one paragraph summary of what was changed and why>"
}}"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )

    raw = message.content[0].text.strip()
    # Strip markdown code fences if present
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  Warning: JSON parse error for {job['title']} — {e}")
        return None


# ── .docx Generation ───────────────────────────────────────────────────────────

def build_docx(tailored: dict, job: dict) -> str:
    """Builds a .docx resume from the tailored JSON and returns the filepath."""

    doc = Document()

    # Page margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Name ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(STELLA_RESUME["name"])
    name_run.bold = True
    name_run.font.size = Pt(16)

    # ── Contact ──
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(STELLA_RESUME["contact"])
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        # Underline via border workaround: add a bottom border paragraph
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F4E79")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_body(text, bold=False, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_bullet(text, size=10):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(size)

    # ── Summary ──
    add_section_heading("Summary")
    add_body(tailored.get("summary", STELLA_RESUME["summary"]))

    # ── Skills ──
    add_section_heading("Skills")
    for category, skills_str in tailored.get("skills", STELLA_RESUME["skills"]).items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        label = p.add_run(f"{category}  ")
        label.bold = True
        label.font.size = Pt(10)
        value = p.add_run(skills_str)
        value.font.size = Pt(10)

    # ── Experience ──
    add_section_heading("Experience")
    for exp in tailored.get("experience", STELLA_RESUME["experience"]):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        title_run = p.add_run(exp["title"])
        title_run.bold = True
        title_run.font.size = Pt(10)

        meta = doc.add_paragraph()
        meta.paragraph_format.space_before = Pt(0)
        meta.paragraph_format.space_after = Pt(1)
        meta_run = meta.add_run(f"{exp['company']}, {exp['location']}  |  {exp['dates']}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for bullet in exp.get("bullets", []):
            add_bullet(bullet)

    # ── Projects ──
    add_section_heading("Projects")
    for proj in STELLA_RESUME["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        title_run = p.add_run(proj["title"])
        title_run.bold = True
        title_run.font.size = Pt(10)
        date_run = p.add_run(f"  {proj['date']}")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        desc = doc.add_paragraph()
        desc.paragraph_format.space_before = Pt(0)
        desc.paragraph_format.space_after = Pt(1)
        desc.add_run(proj["description"]).font.size = Pt(10)

    # ── Education ──
    add_section_heading("Education & Certifications")
    for item in STELLA_RESUME["education"]:
        add_bullet(item)

    # Save file
    safe_company = re.sub(r"[^\w\s-]", "", job["company"]).strip().replace(" ", "_")
    safe_title = re.sub(r"[^\w\s-]", "", job["title"]).strip().replace(" ", "_")
    filename = f"Stella_Fuentes_{safe_title}_{safe_company}.docx"
    doc.save(filename)
    print(f"  Saved: {filename}")
    return filename


# ── Main ───────────────────────────────────────────────────────────────────────

def run(scored_jobs: list[dict]) -> list[dict]:
    """
    Takes scored jobs from search_agent.py.
    Tailors and saves a .docx for each job above MATCH_THRESHOLD.
    Returns list of dicts with job info + filename + changes_made.
    """
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    top_jobs = [j for j in scored_jobs if j.get("match_score", 0) >= MATCH_THRESHOLD]

    print(f"\n=== Resume Tailor: {len(top_jobs)} jobs above {MATCH_THRESHOLD}% threshold ===\n")

    results = []
    for job in top_jobs:
        print(f"Tailoring for: {job['title']} @ {job['company']} ({job['match_score']}%)")
        tailored = tailor_with_claude(job, client)
        if tailored:
            filename = build_docx(tailored, job)
            results.append({
                "title": job["title"],
                "company": job["company"],
                "match_score": job["match_score"],
                "link": job.get("link", ""),
                "source": job.get("source", ""),
                "filename": filename,
                "changes_made": tailored.get("changes_made", ""),
            })

    return results


# Standalone usage — loads from latest job_results JSON
if __name__ == "__main__":
    import glob
    files = sorted(glob.glob("job_results_*.json"), reverse=True)
    if not files:
        print("No job_results JSON found. Run search_agent.py first.")
    else:
        with open(files[0]) as f:
            scored_jobs = json.load(f)
        results = run(scored_jobs)
        print(f"\nDone. {len(results)} resumes generated.")
